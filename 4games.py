import sys
import requests as rq
from bs4 import BeautifulSoup
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm
import progressbar
import re
from urllib import request


imgfolder = "C:\\Users\\2200432\\Desktop\\python\\4gamers\\img\\"
news_type = ["遊戲資訊", "硬體新聞", "影劇動漫", "玩具周邊", "電子競技", "深度專題", "實況直播", "展覽活動"]
game_news = []
hard_news = []
animate_news = []
toys_news = []
gaming_news = []
deep_project = []
live_news = []
event_news = []
array = [game_news, hard_news, animate_news, toys_news, gaming_news, deep_project, live_news, event_news]

def get_url(news_url, user_date):
	r = rq.get(news_url)
	if r.status_code == rq.codes.ok:
		soup = BeautifulSoup(r.text, "html.parser")
		news_url = soup.find_all("h4")
		for a in news_url:
			url = a.select_one("a").get("href")
			data_in_news(url, user_date)

def data_in_news(url, user_date):
	final_text = ""
	r = rq.get(url)
	if r.status_code == rq.codes.ok:
		soup = BeautifulSoup(r.text, "html.parser")
		# time
		time_in_data = soup.find("time")
		date = time_in_data.getText().split()[0]
		date_for_count = int(date[0:4] + date[5:7] + date[8:10])
		date_check(date_for_count, user_date)
		# type
		type_in_data = soup.find("p", class_="category").getText().strip()
		# title
		title_in_data = soup.find("h1").getText().strip()
		# img
		img_url = rq.get(soup.find("img").get("src")).content
		# detail
		p_in_data = soup.find_all("p", class_="", limit=10)
		for word in p_in_data:
			text = word.getText().strip().replace(u'\xa0', u'')
			final_text += text
		check_type(date, type_in_data, title_in_data, url, img_url, final_text)
	else:
		print("Your url is wrong or somrthing else")
		

def date_check(post_time, user_date):
	if user_date > post_time:
		print("資料整合中.........\n")
		docx_maker(user_date)
		print("Done!!")
		sys.exit()

def page_counter(user_date):
	print("\n下載資料中.........\n")
	news_url = "https://www.4gamers.com.tw/news"
	page_count = "?page="
	for i in range(1,3):
		if i == 1:
			get_url(news_url, user_date)
		else:
			get_url(news_url + page_count + str(i), user_date)

def check_type(date, type_in_data, title_in_data, url, img_url, text):
	if type_in_data == news_type[0]:
		game_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 0)
	elif type_in_data == news_type[1]:
		hard_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 1)
	elif type_in_data == news_type[2]:
		animate_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 2)
	elif type_in_data == news_type[3]:
		toys_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 3)
	elif type_in_data == news_type[4]:
		gaming_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 4)
	elif type_in_data == news_type[5]:
		deep_project.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 5)
	elif type_in_data == news_type[6]:
		live_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 6)
	elif type_in_data == news_type[7]:
		event_news.append([date, type_in_data, title_in_data, url, img_url, text])
		img_download(img_url, type_in_data, 7)

def img_download(url, type_in_data, num):
	with open('%s%s\\%s.jpg' % (imgfolder, str(type_in_data), len(array[num])), 'wb') as f:
		f.write(url)

def docx_maker(user_date):
	doc.add_heading('4Games_News', level=0)
	for ary in array:
		data_to_maker(ary)
	doc.save('4Games_news_after_' + str(user_date) + '.docx')

def data_to_maker(array):
	j = 1
	if array != []:
		for i in range(0, len(array)):
			doc.add_heading(array[i][1] + "---第" + str(i + 1) + "則", level=1)
			doc.add_heading(array[i][2], level=2)
			# doc.add_picture(imgfolder + "\\" + str(i) + ".jpg", width=Cm(10))
			doc.add_heading("URL", level=2)
			hyl = doc.add_paragraph(array[i][3] + "\n\n")
			add_hyperlink(hyl, '網站超連結', array[i][3])
			doc.add_heading("日期", level=2)
			doc.add_paragraph(array[i][0])
			doc.add_picture(imgfolder + str(array[i][1]) + "\\" + str(j) + ".jpg", width=Cm(10))
			j += 1
			doc.add_heading('內文', level=2)
			doc.add_paragraph("\n" + array[i][5])
			doc.add_page_break()

def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        pass
 
    try:
        import unicodedata
        unicodedata.numeric(s)
        return True
    except (TypeError, ValueError):
        pass
 
    return False

def add_hyperlink(paragraph, text, url):
       part = paragraph.part
       r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

       hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
       hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

       new_run = docx.oxml.shared.OxmlElement('w:r')
       rPr = docx.oxml.shared.OxmlElement('w:rPr')

       new_run.append(rPr)
       new_run.text = text
       hyperlink.append(new_run)

       r = paragraph.add_run ()
       r._r.append (hyperlink)

       r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
       r.font.underline = True

       return hyperlink

if __name__	== '__main__':
	doc = docx.Document()
	user_date = None
	print("輸入00000000停止程式!")
	while user_date != 00000000:
		user_date = input("輸入終止日(ex : 20201221 ) : ")
		if len(user_date) != 8 or is_number(user_date) == False:
			print("請依照8位數格式輸入(ex : 20201221 ) ") 
		else:
			page_counter(int(user_date))